from docx import Document
import pandas as pd

permanent_document = Document('Dental Injuries summary updated 2021 permanent.docx')
primary_document = Document('Dental Injuries summary updated 2021 primary.docx')


def parse_injuries_tables(document):
    tables = []
    for table in document.tables:
        tab = [['' for i in range(len(table.columns))] for j in range(len(table.rows))]
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if cell.text:
                    tab[i][j] = cell.text.strip().replace('\n', ' ')
        df = pd.DataFrame(tab)
        # Set first row as column names
        df = df.rename(columns=df.iloc[0]).drop(df.index[0])
        # Set first column as row names
        df = df.set_index(df.columns[0])
        tables.append(df)
    # Only first two tables are required in document
    return pd.concat([tables[0], tables[1]], axis=1)


PERMANENT_DF = parse_injuries_tables(permanent_document)
PRIMARY_DF = parse_injuries_tables(primary_document)
